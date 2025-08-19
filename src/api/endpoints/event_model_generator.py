# src/api/endpoints/event_model_generator.py
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
import io
from docx import Document

from src.api.deps import get_db
from src.db import models

router = APIRouter()

def format_event_date_for_doc(evento: models.Evento) -> str:
    """Função auxiliar para formatar a data e hora para o documento."""
    # Adiciona um fuso horário de -04:00 para garantir a data correta
    start_date = evento.data_inicio.strftime('%d/%m/%Y')
    end_date = evento.data_fim.strftime('%d/%m/%Y') if evento.data_fim else None
    
    date_str = start_date
    if end_date and end_date != start_date:
        date_str = f"de {start_date} a {end_date}"
        
    if evento.horario:
        date_str += f", {evento.horario}"
        
    return date_str

@router.get("/", response_class=StreamingResponse)
def get_dynamic_authorization_model(
    evento_id: int,
    db: Session = Depends(get_db)
):
    evento = db.query(models.Evento).filter(models.Evento.id == evento_id).first()
    if not evento:
        raise HTTPException(status_code=404, detail="Evento não encontrado")

    document = Document()
    document.add_heading('AUTORIZAÇÃO PARA PARTICIPAÇÃO EM EVENTO', level=1)
    
    document.add_paragraph(
        'Eu, __________________________________________________, portador(a) do RG nº ____________________ '
        'e CPF nº ____________________, na qualidade de responsável legal pelo(a) aluno(a) '
        '__________________________________________________, autorizo sua participação no seguinte evento:'
    )
    
    table = document.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # --- CORREÇÃO AQUI ---
    formatted_date = format_event_date_for_doc(evento)
    year = evento.data_inicio.year
    
    cells = table.rows
    cells[0].cells[0].text = 'Nome do Evento'
    cells[0].cells[1].text = evento.titulo
    cells[1].cells[0].text = 'Data e Horário'
    cells[1].cells[1].text = formatted_date # Usa a data formatada
    cells[2].cells[0].text = 'Local'
    cells[2].cells[1].text = evento.local_evento or 'A ser definido'

    document.add_paragraph(
        '\nDeclaro estar ciente de todos os detalhes e assumo a responsabilidade por quaisquer '
        'eventualidades. Em caso de emergência, contatar: _________________________.'
    )
    document.add_paragraph('\n\n__________________________________\nAssinatura do Responsável')
    document.add_paragraph(f'Data: ____/____/{year}') # Usa o ano da data de início
    # --- FIM DA CORREÇÃO ---

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    filename = f"autorizacao_{evento.titulo.replace(' ', '_')}.docx"
    return StreamingResponse(
        file_stream,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'}
    )